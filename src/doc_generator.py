from docx import Document

def generate_minutes(optimized_text, output_path):
    doc = Document()
    doc.add_heading("会議記録", level=0)
    doc.add_paragraph("日付: 2025年3月21日")
    doc.add_paragraph("参加者: 山田太郎、佐藤花子、中村一郎")
    
    for line in optimized_text.split("\n"):
        if line.startswith("###"):
            doc.add_heading(line.replace("###", "").strip(), level=1)
        else:
            doc.add_paragraph(line.strip())
    
    doc.save(output_path)
    return output_path